# -*- coding: utf-8 -*-
import sys
import os

TEXT_EXTENSIONS = {'.txt', '.docx', '.doc', '.pdf', '.md', '.csv', '.json', '.xml', '.html', '.htm', 
                   '.log', '.ini', '.cfg', '.conf', '.yaml', '.yml', 
                   '.py', '.c', '.cpp', '.h', '.java', '.js', '.ts', '.css', '.sql', '.sh', '.bat'}

def extract_text(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext not in TEXT_EXTENSIONS:
        return None
    try:
        if ext == '.txt' or ext in ('.md', '.csv', '.json', '.xml', '.html', '.htm', '.log', '.ini', '.cfg', '.conf', '.yaml', '.yml', '.py', '.c', '.cpp', '.h', '.java', '.js', '.ts', '.css', '.sql', '.sh', '.bat'):
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        elif ext == '.docx' or ext == '.doc':
            try:
                from docx import Document
                doc = Document(filepath)
                paragraphs = []
                for p in doc.paragraphs:
                    text = p.text.strip()
                    if text: paragraphs.append(text)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text = cell.text.strip()
                            if text: paragraphs.append(text)
                return '\n'.join(paragraphs)
            except ImportError:
                return ""
        elif ext == '.pdf':
            try:
                try:
                    import pdfplumber
                    with pdfplumber.open(filepath) as pdf:
                        pages = [page.extract_text() for page in pdf.pages if page.extract_text()]
                    return '\n'.join(pages)
                except ImportError:
                    import PyPDF2
                    with open(filepath, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        pages = [page.extract_text() for page in reader.pages if page.extract_text()]
                    return '\n'.join(pages)
            except ImportError:
                return ""
        else:
            return None
    except Exception as e:
        return ""

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python extract_text.py <文件路径>")
        sys.exit(1)
    filepath = sys.argv[1]
    result = extract_text(filepath)
    if result is None:
        print("")
    else:
        print(result)
